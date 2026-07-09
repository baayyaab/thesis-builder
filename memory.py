import streamlit as st
import io
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from thesis_core import (
    ThesisSchema, CommitteeMember, ThesisChapter,
    new_thesis, check_cover_type,
    generate_title_page_en,
    parse_scientific_name,
    monitor_foreign_word_parentheses,
)

# ══════════════════════════════════════════════════════
#  PAGE CONFIG & SESSION STATE
# ══════════════════════════════════════════════════════
st.set_page_config(
    page_title="Modular Thesis Builder — มธ.",
    page_icon="🎓",
    layout="wide",
)
logo_filename = "tu_logo.png"

if "thesis" not in st.session_state:
    st.session_state.thesis = new_thesis("โท")

if "seen_words" not in st.session_state:
    st.session_state.seen_words = set()

if "edu_history" not in st.session_state: st.session_state.edu_history = ""
if "scholarship" not in st.session_state: st.session_state.scholarship = ""
if "publications" not in st.session_state: st.session_state.publications = ""
if "references_raw" not in st.session_state: st.session_state.references_raw = ""

thesis: ThesisSchema = st.session_state.thesis

# ══════════════════════════════════════════════════════
#  HELPER FUNCTIONS FOR AUTOMATIC PAGE NUMBERING IN (X)
# ══════════════════════════════════════════════════════

def add_thesis_heading(doc, text, level=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(18 if level == 0 else 16)
    run.bold = True
    
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
    else:
        # ระยะเยื้องเริ่มต้นของหัวข้อ
        indent_map = {1: 0.0, 2: 0.8, 3: 1.1, 4: 1.4}
        indent_val = indent_map.get(level, 0)
        
        p.paragraph_format.left_indent = Inches(indent_val)
        
        # ปรับ Hanging Indent: 
        # ทำให้บรรทัดที่ 2 เป็นต้นไป ย้อนกลับมาให้ตรงกับตำแหน่งตัวอักษรแรก
        # เราต้องกำหนด first_line_indent เป็นค่าติดลบของระยะเยื้องที่ต้องการ
        # ในที่นี้สมมติความกว้างของเลขหัวข้อประมาณ 0.4 นิ้ว
        p.paragraph_format.first_line_indent = Inches(-0.4) 
        
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(24)
    return p

def add_custom_para(text="", size=16, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, space_before=0):
    """
    ฟังก์ชันสำหรับเนื้อหาปกติ (ค่าเริ่มต้น bold=False จะไม่เป็นตัวหนา)
    """
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(size)
        run.bold = bold  # จะเป็นตัวหนาเฉพาะเมื่อสั่งให้เป็นเท่านั้น
    return p

def add_tu_style_page_number(run):
    """ฟังก์ชันระดับ Low-level สำหรับแทรก Field Code เลขหน้าอัตโนมัติใน Word ครอบด้วยวงเล็บ ( )"""
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def clean_text(raw: str) -> tuple[str, list, list]:
    sci = parse_scientific_name(raw)
    text_after_sci = sci["processed_text"]
    paren = monitor_foreign_word_parentheses(text_after_sci, st.session_state.seen_words)
    st.session_state.seen_words = paren["seen_words"]
    return paren["processed_text"], sci["found_names"], paren["removed_pairs"]

def member_form(key_prefix: str, label: str, default: CommitteeMember | None = None) -> CommitteeMember:
    d = default or CommitteeMember("", "", "", "", label)
    c1, c2, c3, c4 = st.columns([1, 2, 1, 2])
    with c1: title_th = st.text_input("คำนำหน้า (ไทย)", value=d.title_th, key=f"{key_prefix}_tth")
    with c2: name_th  = st.text_input("ชื่อ-นามสกุล (ไทย)", value=d.name_th, key=f"{key_prefix}_nth")
    with c3: title_en = st.text_input("Title (EN)", value=d.title_en, key=f"{key_prefix}_ten")
    with c4: name_en  = st.text_input("Full Name (EN)", value=d.name_en, key=f"{key_prefix}_nen")
    return CommitteeMember(title_th, name_th, title_en, name_en, label)

# ══════════════════════════════════════════════════════
#  PART 3: PROFESSIONAL DOCX EXPORTER ENGINE (มธ. 2026)
# ══════════════════════════════════════════════════════
def build_docx_file(data: ThesisSchema, citation_style: str) -> io.BytesIO:
    doc = Document()
    
    # ตั้งค่าระยะขอบ (Margins) ตามเกณฑ์ มธ. (บน 1.5 ซ้าย 1.5 ล่าง 1 ขวา 1 นิ้ว)
    for section in doc.sections:
        section.top_margin = Inches(1.5) 
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        
    # ตั้งค่า Normal Style เป็น TH Sarabun New ทั้งเล่ม
    style = doc.styles['Normal']
    font = style.font
    font.name = 'TH Sarabun New'
    font.size = Pt(16)
    
    def add_custom_para(text="", size=16, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, space_before=0):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if text:
            run = p.add_run(text)
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(size)
            run.bold = bold
        return p


    # 🛑 --- PAGE 1: หน้าปกนอก (Thai Outer Cover) ---
    table = doc.add_table(rows=2, cols=1)
    table.autofit = False
    
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(r'<w:trHeight %s w:val="7500" w:hRule="atLeast"/>' % nsdecls('w')))
    
    tblPr = table._tbl.tblPr
    borders = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
    tblPr.append(borders)

    cell_top = table.cell(0, 0)
    p_img = cell_top.paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_after = Pt(12)
    if os.path.exists(logo_filename):
        run_img = p_img.add_run()
        run_img.add_picture(logo_filename, width=Inches(1.0))
    else:
        p_img.add_run("[ กรุณาวางไฟล์ตราธรรมจักรไว้ในโฟลเดอร์ ]")

    p_title = cell_top.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(36)
    run_title = p_title.add_run(data.title_th)
    run_title.font.size = Pt(20)
    run_title.bold = True

    p_by = cell_top.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_by.paragraph_format.space_after = Pt(12)
    run_by = p_by.add_run("โดย")
    run_by.font.size = Pt(18)
    run_by.bold = True

    p_auth = cell_top.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_auth = p_auth.add_run(data.author_th)
    run_auth.font.size = Pt(18)
    run_auth.bold = True

    cell_bottom = table.cell(1, 0)
    p_meta = cell_bottom.paragraphs[0]
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    field_text = data.field_of_study_th.replace("สาขาวิชา", "", 1) if data.field_of_study_th.startswith("สาขาวิชา") else data.field_of_study_th

    cover_meta_text = (
        f"วิทยานิพนธ์นี้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตร\n"
        f"{data.degree_name_th}\n"
        f"สาขาวิชา{field_text}\n"
        f"คณะ{data.faculty_th} มหาวิทยาลัยธรรมศาสตร์\n"
        f"ปีการศึกษา {data.academic_year_be}"
    )
    run_meta = p_meta.add_run(cover_meta_text)
    run_meta.font.size = Pt(18)
    run_meta.bold = True

    # 🛑 --- PAGE 2: หน้าปกในอังกฤษ (English Title Page) ---
    doc.add_page_break()
    doc_label_en = "DISSERTATION" if data.degree_level == "เอก" else ("AN INDEPENDENT STUDY" if data.degree_level == "IS" else "THESIS")

    add_custom_para(data.title_en.upper(), size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_custom_para("BY", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=36, space_after=18)
    add_custom_para(data.author_en.upper(), size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=48)
    
    field_en_text = data.field_of_study_en.upper().replace("FIELD OF STUDY IN ", "", 1).replace("FIELD OF STUDY ", "", 1)
        
    en_meta_block = (
        f"A {doc_label_en} SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS\n"
        f"FOR THE DEGREE OF {data.degree_name_en.upper()}\n"
        f"IN {field_en_text}\n"
        f"FACULTY OF {data.faculty_en.upper()}\n"
        f"THAMMASAT UNIVERSITY\n"
        f"ACADEMIC YEAR {data.academic_year_ce}"
    )
    add_custom_para(en_meta_block, size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=24)

    # 🛑 --- PAGE 3: หน้าอนุมัติ (Approval Page) ---
    doc.add_page_break()
    doc_label_th = "ดุษฎีนิพนธ์" if data.degree_level == "เอก" else ("สารนิพนธ์การค้นคว้าอิสระ" if data.degree_level == "IS" else "วิทยานิพนธ์")

    add_custom_para("มหาวิทยาลัยธรรมศาสตร์", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_custom_para(f"คณะ{data.faculty_th}", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    
    add_custom_para("", size=16, space_after=0)
    add_custom_para(doc_label_th, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    
    add_custom_para("", size=16, space_after=0)
    add_custom_para("ของ", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    
    add_custom_para("", size=16, space_after=0)
    add_custom_para(data.author_th, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    
    add_custom_para("", size=16, space_after=0)
    add_custom_para("เรื่อง", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    
    add_custom_para("", size=16, space_after=0)
    add_custom_para(data.title_th, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
   
    add_custom_para("", size=16, space_after=0)
    clean_field_th = data.field_of_study_th.replace("สาขาวิชา", "").strip()
    full_approval_line = f"ได้รับการตรวจสอบและอนุมัติ ให้เป็นส่วนหนึ่งของการศึกษาตามหลักสูตร{data.degree_name_th} ({clean_field_th})"
    add_custom_para(full_approval_line, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    
    add_custom_para("", size=16, space_after=0)
    add_custom_para(f"เมื่อ วันที่ {data.approval_day} {data.approval_month_th} พ.ศ. {data.approval_year_be}", size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    sign_table = doc.add_table(rows=0, cols=2)
    sign_table.autofit = False
    sign_table.columns[0].width = Inches(3.8)
    sign_table.columns[1].width = Inches(2.7)
    
    tblPr_sign = sign_table._tbl.tblPr
    borders_sign = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
    tblPr_sign.append(borders_sign)

    def add_sign_row(position_text, member: CommitteeMember):
        if not member.name_th.strip(): return
        row = sign_table.add_row()
        p_pos = row.cells[0].paragraphs[0]
        p_pos.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_pos.paragraph_format.space_after = Pt(12)
        p_pos.add_run(position_text).font.size = Pt(16)
        
        p_line = row.cells[1].paragraphs[0]
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_line.paragraph_format.space_after = Pt(12)
        
        run_space = p_line.add_run("\u00A0" * 35)
        run_space.font.size = Pt(16)
        run_space.font.underline = True
        p_line.add_run("\n")
        run_l = p_line.add_run(f"( {member.title_th}{member.name_th} )")
        run_l.font.size = Pt(16)

    add_sign_row("ประธานกรรมการสอบวิทยานิพนธ์", data.committee_chair)
    add_sign_row("กรรมการและอาจารย์ที่ปรึกษา.", data.advisor_main)
    if data.advisor_co:
        for co_adv in data.advisor_co: add_sign_row("กรรมการและอาจารย์ที่ปรึกษา.", co_adv)
    if data.examiners:
        for examiner in data.examiners: add_sign_row("กรรมการสอบวิทยานิพนธ์", examiner)
    add_sign_row("คณบดี", data.dean)

    # 🛑 --- ตั้งค่าเลขหน้าในวงเล็บ (1) สำหรับส่วนเนื้อหาถัดจากนี้เป็นต้นไป ---
    body_section = doc.add_section()
    body_section.header.is_linked_to_previous = False
    
    sectPr = body_section._sectPr
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')  
    sectPr.append(pgNumType)
    
    header_para = body_section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_para.paragraph_format.space_after = Pt(0)
    
    run_bracket_open = header_para.add_run("(")
    run_bracket_open.font.name = 'TH Sarabun New'
    run_bracket_open.font.size = Pt(16)
    
    run_num = header_para.add_run()
    run_num.font.name = 'TH Sarabun New'
    run_num.font.size = Pt(16)
    add_tu_style_page_number(run_num)
    
    run_bracket_close = header_para.add_run(")")
    run_bracket_close.font.name = 'TH Sarabun New'
    run_bracket_close.font.size = Pt(16)

    # 🛑 --- PAGE 4: บทคัดย่อภาษาไทย ---
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.autofit = False
    meta_table.columns[0].width = Inches(3.0)  
    meta_table.columns[1].width = Inches(3.5)  
    
    tblPr_meta = meta_table._tbl.tblPr
    borders_meta = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
    tblPr_meta.append(borders_meta)
    
    clean_field = data.field_of_study_th.replace("สาขาวิชา", "").strip()
    clean_faculty = data.faculty_th.replace("คณะ", "").strip()
    
    abstract_meta_rows = [
        ("หัวข้อวิทยานิพนธ์", data.title_th),
        ("ชื่อผู้เขียน", data.author_th),
        ("ชื่อปริญญา", data.degree_name_th),
        ("สาขาวิชา/คณะ/มหาวิทยาลัย", f"{clean_field}\nวิทยาลัย{clean_faculty} มหาวิทยาลัยธรรมศาสตร์" if "วิทยาลัย" in clean_faculty else f"{clean_field}\nคณะ{clean_faculty}\nมหาวิทยาลัยธรรมศาสตร์"),
        ("อาจารย์ที่ปรึกษาวิทยานิพนธ์", f"{data.advisor_main.title_th}{data.advisor_main.name_th}")
    ]
    if data.advisor_co:
        for co_adv in data.advisor_co:
            if hasattr(co_adv, 'name_th') and co_adv.name_th.strip():
                abstract_meta_rows.append(("อาจารย์ที่ปรึกษาร่วม", f"{co_adv.title_th}{co_adv.name_th}"))
    abstract_meta_rows.append(("ปีการศึกษา", str(data.academic_year_be)))
    
    for title_str, value_str in abstract_meta_rows:
        row_m = meta_table.add_row()
        p_t = row_m.cells[0].paragraphs[0]
        p_t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_t.paragraph_format.space_after = Pt(3)
        run_t = p_t.add_run(title_str)
        run_t.font.name = 'TH Sarabun New'
        run_t.font.size = Pt(16)
        
        p_v = row_m.cells[1].paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_v.paragraph_format.space_after = Pt(3)
        p_v.paragraph_format.line_spacing = 1.15
        run_v = p_v.add_run(value_str)
        run_v.font.name = 'TH Sarabun New'
        run_v.font.size = Pt(16)

    p_abs_head = doc.add_paragraph()
    p_abs_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_abs_head.paragraph_format.space_before = Pt(24)
    p_abs_head.paragraph_format.space_after = Pt(18)
    run_abs_head = p_abs_head.add_run("บทคัดย่อ")
    run_abs_head.font.name = 'TH Sarabun New'
    run_abs_head.font.size = Pt(18)
    run_abs_head.bold = True

    if data.abstract_th.strip():
        paragraphs_th = data.abstract_th.split('\n')
        for para in paragraphs_th:
            if para.strip():
                p_body = add_custom_para(para.strip(), 16, False, space_after=6)
                p_body.paragraph_format.first_line_indent = Inches(0.8)

    kw_th_final = getattr(data, "keywords_th", "") or ""
    if str(kw_th_final).strip():
        add_custom_para("", size=16, space_after=0)  
        p_kw_th = doc.add_paragraph()
        p_kw_th.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_kw_th.paragraph_format.space_after = Pt(0)
        
        run_kw_label_th = p_kw_th.add_run("คำสำคัญ: ")
        run_kw_label_th.font.name = 'TH Sarabun New'
        run_kw_label_th.font.size = Pt(16)
        run_kw_label_th.bold = True
        
        run_kw_val_th = p_kw_th.add_run(str(kw_th_final).strip())
        run_kw_val_th.font.name = 'TH Sarabun New'
        run_kw_val_th.font.size = Pt(16)

    # 🛑 --- PAGE 5: Abstract (English) ---
    doc.add_page_break()
    
    en_meta_table = doc.add_table(rows=0, cols=2)
    en_meta_table.autofit = False
    en_meta_table.columns[0].width = Inches(2.3)  
    en_meta_table.columns[1].width = Inches(4.2)  
    
    tblPr_en_meta = en_meta_table._tbl.tblPr
    borders_en_meta = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
    tblPr_en_meta.append(borders_en_meta)
    
    title_label = "Dissertation Title" if data.degree_level == "เอก" else "Thesis Title"
    
    clean_field_en = data.field_of_study_en.replace("Field of Study in", "").replace("Field of Study", "").strip()
    clean_faculty_en = data.faculty_en.replace("Faculty of", "").strip()
    
    abstract_en_meta_rows = [
        (title_label, data.title_en.upper()), 
        ("Author", data.author_en),
        ("Degree", data.degree_name_en),
        ("Department/Faculty/University", f"{clean_field_en}\nFaculty of {clean_faculty_en}\nThammasat University"),
        ("Thesis Advisor", f"{data.advisor_main.title_en} {data.advisor_main.name_en}" if data.advisor_main.title_en else data.advisor_main.name_en)
    ]
    if data.advisor_co:
        for co_adv in data.advisor_co:
            if hasattr(co_adv, 'name_en') and co_adv.name_en.strip():
                adv_title = f"{co_adv.title_en} " if co_adv.title_en else ""
                abstract_en_meta_rows.append(("Thesis Co-Advisor", f"{adv_title}{co_adv.name_en}"))
    abstract_en_meta_rows.append(("Academic Year", str(data.academic_year_ce)))
    
    for t_str, v_str in abstract_en_meta_rows:
        row_en = en_meta_table.add_row()
        p_en_t = row_en.cells[0].paragraphs[0]
        p_en_t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_en_t.paragraph_format.space_after = Pt(3)
        run_en_t = p_en_t.add_run(t_str)
        run_en_t.font.name = 'TH Sarabun New'
        run_en_t.font.size = Pt(16)
        
        p_en_v = row_en.cells[1].paragraphs[0]
        p_en_v.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_en_v.paragraph_format.space_after = Pt(3)
        p_en_v.paragraph_format.line_spacing = 1.15
        run_en_v = p_en_v.add_run(v_str)
        run_en_v.font.name = 'TH Sarabun New'
        run_en_v.font.size = Pt(16)

    add_custom_para("", size=16, space_after=0)
    
    p_en_abs_head = doc.add_paragraph()
    p_en_abs_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_en_abs_head.paragraph_format.space_before = Pt(6)
    p_en_abs_head.paragraph_format.space_after = Pt(6)
    run_en_abs_head = p_en_abs_head.add_run("ABSTRACT")
    run_en_abs_head.font.name = 'TH Sarabun New'
    run_en_abs_head.font.size = Pt(18)
    run_en_abs_head.bold = True

    add_custom_para("", size=16, space_after=0)

    if data.abstract_en.strip():
        paragraphs_en = data.abstract_en.split('\n')
        for para in paragraphs_en:
            if para.strip():
                p_body_en = add_custom_para(para.strip(), 16, False, space_after=6)
                p_body_en.paragraph_format.first_line_indent = Inches(0.8) 

    kw_en_final = getattr(data, "keywords_en", "") or ""
    if str(kw_en_final).strip():
        add_custom_para("", size=16, space_after=0)  
        p_kw_en = doc.add_paragraph()
        p_kw_en.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_kw_en.paragraph_format.space_after = Pt(0)
        
        run_kw_label_en = p_kw_en.add_run("Keywords: ")
        run_kw_label_en.font.name = 'TH Sarabun New'
        run_kw_label_en.font.size = Pt(16)
        run_kw_label_en.bold = True
        
        run_kw_val_en = p_kw_en.add_run(str(kw_en_final).strip())
        run_kw_val_en.font.name = 'TH Sarabun New'
        run_kw_val_en.font.size = Pt(16)

# 🛑 --- PAGE 6: กิตติกรรมประกาศ ---
    if data.acknowledgement.strip():
        doc.add_page_break()
        
        p_ack_head = doc.add_paragraph()
        p_ack_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ack_head.paragraph_format.space_before = Pt(0)
        p_ack_head.paragraph_format.space_after = Pt(0)
        run_ack_head = p_ack_head.add_run("กิตติกรรมประกาศ")
        run_ack_head.font.name = 'TH Sarabun New'
        run_ack_head.font.size = Pt(18)
        run_ack_head.bold = True

        add_custom_para("", size=16, space_after=0)

        paragraphs_ack = data.acknowledgement.split('\n')
        for para in paragraphs_ack:
            if para.strip():
                p_body_ack = add_custom_para(para.strip(), 16, False, space_after=6)
                p_body_ack.paragraph_format.first_line_indent = Inches(0.8)

        add_custom_para("", size=16, space_after=0)
        add_custom_para("", size=16, space_after=0)

        p_author_name = doc.add_paragraph()
        p_author_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_author_name.paragraph_format.right_indent = Inches(0.5)
        p_author_name.paragraph_format.space_after = Pt(0)
        
        run_auth_name = p_author_name.add_run(data.author_th)
        run_auth_name.font.name = 'TH Sarabun New'
        run_auth_name.font.size = Pt(16)

    # 🛑 --- ✨ หน้าสารบัญที่ปรับปรุงใหม่ (แทรกแถวว่างเพื่อเว้น 1 บรรทัดเต็มชัวร์ ๆ) ✨ ---
    doc.add_page_break()
    
    # 1. หัวข้อสารบัญ ขนาด 18 ตัวหนา จัดกึ่งกลางหน้ากระดาษ
    add_custom_para(text="สารบัญ", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=0)
    
    # 2. สร้างโครงตารางจัดหน้าข้อความชิดซ้าย-ขวา
    toc_table = doc.add_table(rows=0, cols=2)
    toc_table.autofit = False
    toc_table.columns[0].width = Inches(5.0)
    toc_table.columns[1].width = Inches(1.5)
    
    # ซ่อนขอบตาราง เพื่อความสะอาดและถูกต้องตามฟอร์แมต มธ.
    tblPr_toc = toc_table._tbl.tblPr
    borders_toc = parse_xml(r'<w:tblBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/><w:insideH w:val="none"/><w:insideV w:val="none"/></w:tblBorders>' % nsdecls('w'))
    tblPr_toc.append(borders_toc)
    
    # ฟังก์ชันช่วยเพิ่มแถวเปล่า (เว้น 1 บรรทัดเต็ม)
    def add_blank_row():
        row_blank = toc_table.add_row()
        # คอลัมน์ซ้าย
        p_b1 = row_blank.cells[0].paragraphs[0]
        p_b1.paragraph_format.space_before = Pt(0)
        p_b1.paragraph_format.space_after = Pt(0)
        r_b1 = p_b1.add_run("")  # ปล่อยว่าง
        r_b1.font.name = 'TH Sarabun New'
        r_b1.font.size = Pt(16)
        # คอลัมน์ขวา
        p_b2 = row_blank.cells[1].paragraphs[0]
        p_b2.paragraph_format.space_before = Pt(0)
        p_b2.paragraph_format.space_after = Pt(0)

    # บรรทัดถัดจากคำว่า "สารบัญ" -> เว้น 1 บรรทัดเต็ม
    add_blank_row()
    
    # ส่วนขวาบน: ใส่คำว่า "หน้า" เป็นตัวปกติ
    row_head = toc_table.add_row()
    p_h1 = row_head.cells[0].paragraphs[0]
    p_h1.paragraph_format.space_before = Pt(0)
    p_h1.paragraph_format.space_after = Pt(0)
    
    p_h2 = row_head.cells[1].paragraphs[0]
    p_h2.paragraph_format.space_before = Pt(0)
    p_h2.paragraph_format.space_after = Pt(0)
    p_h2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_h2 = p_h2.add_run("หน้า")
    run_h2.font.name = 'TH Sarabun New'
    run_h2.font.size = Pt(16)
    
    # รายการหัวข้อด้านซ้าย
    toc_items = [
        "บทคัดย่อภาษาไทย",
        "บทคัดย่อภาษาอังกฤษ",
        "กิตติกรรมประกาศ",
        "สารบัญตาราง",
        "สารบัญภาพ"
    ]
    
    for i, item in enumerate(toc_items):
        # ทุก ๆ ช่องว่างระหว่างหัวข้อ จะถูกคั่นด้วยแถวว่าง 1 บรรทัดเต็ม (ยกเว้นก่อนหน้าบทคัดย่อภาษาไทยที่มีคำว่าหน้าอยู่แล้ว)
        if i > 0:
            add_blank_row()
            
        row_item = toc_table.add_row()
        p_left = row_item.cells[0].paragraphs[0]
        p_left.paragraph_format.space_before = Pt(0)
        p_left.paragraph_format.space_after = Pt(0)
        run_item = p_left.add_run(item)
        run_item.font.name = 'TH Sarabun New'
        run_item.font.size = Pt(16)
        
        # คอลัมน์ขวาปล่อยว่างตามเงื่อนไข
        p_right = row_item.cells[1].paragraphs[0]
        p_right.paragraph_format.space_before = Pt(0)
        p_right.paragraph_format.space_after = Pt(0)

   # 🛑 --- PAGES 7+: เนื้อหาแต่ละบท ---
    for ch in data.chapters:
        doc.add_page_break()
        # ชื่อบท (ตัวหนาตามที่ตกลง)
        add_thesis_heading(doc, f"บทที่ {ch.chapter_number}", level=0)
        add_thesis_heading(doc, ch.title_th, level=0)
        
        if ch.body.strip():
            # ใช้การวนลูปแบบตรวจสอบเงื่อนไขที่รัดกุมขึ้น
            for line in ch.body.split('\n'):
                line = line.strip()
                if not line: continue
                
                # ตรวจสอบว่าเป็นหัวข้อหรือไม่ (ต้องขึ้นต้นด้วยตัวเลข หรือ (1))
                # วิธีนี้ช่วยแยกหัวข้อออกจากเนื้อหาได้แม่นยำขึ้น
                is_heading = False
                level = 0
                
                if line.startswith('('): 
                    is_heading = True; level = 4
                elif line.count('.') == 3 and line[0].isdigit():
                    is_heading = True; level = 3
                elif line.count('.') == 2 and line[0].isdigit():
                    is_heading = True; level = 2
                elif line.count('.') == 1 and line[0].isdigit():
                    is_heading = True; level = 1
                
                if is_heading:
                    # หัวข้อ: ใช้ฟังก์ชัน add_thesis_heading (จะเป็นตัวหนาอัตโนมัติ)
                    add_thesis_heading(doc, line, level=level)
                else:
                    # เนื้อหาปกติ: ใช้ add_custom_para โดยบังคับ bold=False ชัดเจน
                    p_body = add_custom_para(line, 16, bold=False, space_after=6)
                    p_body.paragraph_format.first_line_indent = Inches(0.8)

    # 🛑 --- รายการอ้างอิง และส่วนท้ายเล่ม ---
    # (ใช้ add_custom_para ปกติที่คุณมีอยู่ได้เลยครับ)
    if st.session_state.references_raw.strip():
        doc.add_page_break()
        add_custom_para("รายการอ้างอิง", 18, True, WD_ALIGN_PARAGRAPH.CENTER, 24)
        for r in st.session_state.references_raw.split('\n'):
            if r.strip():
                p_ref = add_custom_para(r.strip(), 16, False, space_after=12)
                p_ref.paragraph_format.left_indent = Inches(0.5)
                p_ref.paragraph_format.first_line_indent = Inches(-0.5)

    if data.appendix.strip():
        doc.add_page_break()
        add_custom_para("ภาคผนวก", 18, True, WD_ALIGN_PARAGRAPH.CENTER, 24)
        p_app = add_custom_para(data.appendix, 16, False)
        p_app.paragraph_format.first_line_indent = Inches(0.8)

    if data.author_biography.strip():
        doc.add_page_break()
        add_custom_para("ประวัติผู้เขียน", 18, True, WD_ALIGN_PARAGRAPH.CENTER, 24)
        for b in data.author_biography.split('\n\n'):
            if b.strip(): add_custom_para(b.strip(), 16, False, space_after=12)
    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    return target

# ══════════════════════════════════════════════════════
#  SIDEBAR MANAGEMENT
# ══════════════════════════════════════════════════════
with st.sidebar:
    st.image("https://upload.wikimedia.org/wikipedia/th/thumb/4/4e/Thammasat_University_Logo.svg/200px-Thammasat_University_Logo.svg.png", width=80)
    st.title("⚙️ ตั้งค่าเล่ม")
    st.divider()

    degree_level = st.selectbox("🎓 ระดับปริญญา", ["โท", "IS", "เอก"], index=["โท", "IS", "เอก"].index(thesis.degree_level))
    if degree_level != thesis.degree_level:
        thesis.degree_level = degree_level

    main_lang = st.selectbox("🌐 ภาษาหลักของเล่ม", ["ภาษาไทย", "ภาษาอังกฤษ"])
    citation_style = st.selectbox("📚 รูปแบบการอ้างอิง", ["APA", "Turabian", "Vancouver"])

    st.divider()
    cover = check_cover_type(degree_level)
    st.markdown("**🎨 สีปกที่ระบบเลือกให้**")
    st.info(f"พื้นปก: **{cover['bg_color']}** \nตัวอักษร: **{cover['text_color']}** \nประเภทเล่ม: **{cover['doc_label']}**")

    st.divider()
    if st.button("🗑️ รีเซ็ตเล่มใหม่", use_container_width=True):
        st.session_state.thesis = new_thesis(degree_level)
        st.session_state.seen_words = set()
        st.session_state.edu_history = ""
        st.session_state.scholarship = ""
        st.session_state.publications = ""
        st.session_state.references_raw = ""
        st.rerun()

# ══════════════════════════════════════════════════════
#  MAIN APP INTERFACE
# ══════════════════════════════════════════════════════
st.title("🎓 Modular Thesis Builder — มหาวิทยาลัยธรรมศาสตร์")
st.caption(f"ระดับ: {degree_level} · ภาษา: {main_lang} · อ้างอิง: {citation_style}")
st.divider()

tab1, tab2, tab3, tab4 = st.tabs([
    "📋 หน้าปก & กรรมการ",
    "📄 ส่วนนำเรื่อง",
    "📝 เนื้อหาแต่ละบท",
    "📎 ส่วนท้าย & ประวัติ",
])

with tab1:
    st.subheader("📋 Metadata — ข้อมูลหน้าปก")
    col_th, col_en = st.columns(2)
    with col_th:
        thesis.title_th = st.text_input("ชื่อเรื่อง (ภาษาไทย)", value=thesis.title_th)
        thesis.author_th = st.text_input("ชื่อผู้เขียน (ภาษาไทย) — รวมคำนำหน้า", value=thesis.author_th)
        thesis.degree_name_th = st.text_input("ชื่อเต็มปริญญา (ไทย)", value=thesis.degree_name_th)
        thesis.field_of_study_th = st.text_input("สาขาวิชา (ไทย)", value=thesis.field_of_study_th)
        thesis.faculty_th = st.text_input("คณะ (ไทย)", value=thesis.faculty_th)
    with col_en:
        thesis.title_en = st.text_input("Title (English)", value=thesis.title_en)
        thesis.author_en = st.text_input("Author (English)", value=thesis.author_en)
        thesis.degree_name_en = st.text_input("Degree Name (EN)", value=thesis.degree_name_en)
        thesis.field_of_study_en = st.text_input("Field of Study (EN)", value=thesis.field_of_study_en)
        thesis.faculty_en = st.text_input("Faculty (EN)", value=thesis.faculty_en)

    st.divider()
    col_y1, col_y2, col_d, col_m, col_my = st.columns(5)
    with col_y1:
        thesis.academic_year_be = st.number_input("ปีการศึกษา (พ.ศ.)", min_value=2500, max_value=2600, value=thesis.academic_year_be or 2567, step=1)
    with col_y2:
        st.text_input("ปีการศึกษา (ค.ศ.)", value=str(thesis.academic_year_be - 543), disabled=True)
        thesis.academic_year_ce = thesis.academic_year_be - 543
    with col_d:
        thesis.approval_day = st.number_input("วันที่อนุมัติ", min_value=1, max_value=31, value=thesis.approval_day or 1, step=1)
    with col_m:
        months_th = ["มกราคม","กุมภาพันธ์","มีนาคม","เมษายน","พฤษภาคม","มิถุนายน","กรกฎาคม","สิงหาคม","กันยายน","ตุลาคม","พฤศจิกายน","ธันวาคม"]
        idx_m = months_th.index(thesis.approval_month_th) if thesis.approval_month_th in months_th else 0
        thesis.approval_month_th = st.selectbox("เดือนอนุมัติ (ไทย)", months_th, index=idx_m)
    with col_my:
        thesis.approval_year_be = st.number_input("ปีอนุมัติ (พ.ศ.)", min_value=2500, max_value=2600, value=thesis.approval_year_be or 2568, step=1)

    st.divider()
    st.subheader("👥 คณะกรรมการอนุมัติ")
    committee_tabs = st.tabs(["👤 ประธานกรรมการสอบ", "👨‍🏫 กรรมการและอาจารย์ที่ปรึกษา", "📝 กรรมการสอบ", "🏛️ คณบดี"])
    
    with committee_tabs[0]: thesis.committee_chair = member_form("chair", "ประธานกรรมการ", thesis.committee_chair)
    with committee_tabs[1]:
        thesis.advisor_main = member_form("adv_main", "อาจารย์ที่ปรึกษาหลัก", thesis.advisor_main)
        st.divider()
        has_co_advisor = st.checkbox("มีอาจารย์ที่ปรึกษาร่วม", value=len(thesis.advisor_co) > 0)
        if has_co_advisor:
            if not thesis.advisor_co: thesis.advisor_co = [CommitteeMember("","","","","อาจารย์ที่ปรึกษาร่วม")]
            for i, co in enumerate(thesis.advisor_co): thesis.advisor_co[i] = member_form(f"co_{i}", f"ที่ปรึกษาร่วม #{i+1}", co)
            if st.button("➕ เพิ่มที่ปรึกษาร่วม"):
                thesis.advisor_co.append(CommitteeMember("","","","","อาจารย์ที่ปรึกษาร่วม"))
                st.rerun()
        else: thesis.advisor_co = []
    with committee_tabs[2]:
        if not thesis.examiners: thesis.examiners = [CommitteeMember("","","","","กรรมการสอบ")]
        for i, ex in enumerate(thesis.examiners):
            col_ex, col_del = st.columns([10, 1])
            with col_ex: thesis.examiners[i] = member_form(f"exam_{i}", f"กรรมการสอบ #{i+1}", ex)
            with col_del:
                st.markdown("<br>", unsafe_allow_html=True)
                if len(thesis.examiners) > 1 and st.button("🗑️", key=f"del_ex_{i}"):
                    thesis.examiners.pop(i)
                    st.rerun()
        if st.button("➕ เพิ่มกรรมการสอบ"):
            thesis.examiners.append(CommitteeMember("","","","","กรรมการสอบ"))
            st.rerun()
    with committee_tabs[3]: thesis.dean = member_form("dean", "คณบดี", thesis.dean)

with tab2:
    st.subheader("📄 ส่วนนำเรื่อง (Front Matter)")
    thesis.abstract_th = st.text_area("บทคัดย่อภาษาไทย", value=thesis.abstract_th, height=200)
    
    init_kw_th = getattr(thesis, "keywords_th", "") or ""
    thesis.keywords_th = st.text_input("คำสำคัญ (ภาษาไทย)", value=str(init_kw_th))
    
    st.divider()
    thesis.abstract_en = st.text_area("Abstract (English)", value=thesis.abstract_en, height=200)
    
    init_kw_en = getattr(thesis, "keywords_en", "") or ""
    thesis.keywords_en = st.text_input("Keywords (English)", value=str(init_kw_en))
    
    st.divider()
    thesis.acknowledgement = st.text_area("กิตติกรรมประกาศ", value=thesis.acknowledgement, height=150)

    if st.button("✨ คลิกเพื่อขัดเกลาข้อความส่วนนำ", type="secondary", use_container_width=True):
        th_clean, _, _ = clean_text(thesis.abstract_th)
        en_clean, _, _ = clean_text(thesis.abstract_en)
        ack_clean, _, _ = clean_text(thesis.acknowledgement)
        thesis.abstract_th = th_clean
        thesis.abstract_en = en_clean
        thesis.acknowledgement = ack_clean
        st.rerun()

with tab3:
    st.subheader("📝 เนื้อหาแต่ละบท")
    for i, ch in enumerate(thesis.chapters):
        with st.expander(f"📖 บทที่ {ch.chapter_number} — {ch.title_th or '(ยังไม่ตั้งชื่อ)'}", expanded=(i == 0)):
            col_cth, col_cen = st.columns(2)
            with col_cth: thesis.chapters[i].title_th = st.text_input("ชื่อบท (ไทย)", value=ch.title_th, key=f"ch_tth_{i}")
            with col_cen: thesis.chapters[i].title_en = st.text_input("Chapter Title (EN)", value=ch.title_en, key=f"ch_ten_{i}")
            thesis.chapters[i].body = st.text_area("เนื้อหาบท", value=ch.body, height=250, key=f"ch_body_{i}", label_visibility="collapsed")

            if st.button(f"✨ ขัดเกลาคำอัตโนมัติใน บทที่ {ch.chapter_number}", key=f"scan_btn_{i}"):
                cleaned, _, _ = clean_text(thesis.chapters[i].body)
                thesis.chapters[i].body = cleaned
                st.rerun()

    if st.button("➕ เพิ่มบทใหม่", type="primary", use_container_width=True):
        thesis.add_chapter("(ยังไม่ตั้งชื่อ)", "")
        st.rerun()

with tab4:
    st.subheader("📎 รายการอ้างอิง & ภาคผนวก")
    st.session_state.references_raw = st.text_area("รายการอ้างอิง (ขึ้นบรรทัดใหม่ต่อ 1 รายการ)", value=st.session_state.references_raw, height=200)
    thesis.appendix = st.text_area("ภาคผนวก", value=thesis.appendix, height=150)
    st.divider()
    st.subheader("👤 ประวัติผู้เขียน")
    st.session_state.edu_history = st.text_area("🎓 วุฒิการศึกษา", value=st.session_state.edu_history)
    st.session_state.scholarship = st.text_area("🏅 ทุนการศึกษา", value=st.session_state.scholarship)
    st.session_state.publications = st.text_area("📰 ผลงานทางวิชาการ", value=st.session_state.publications)

    bio_parts = []
    if st.session_state.edu_history.strip(): bio_parts.append(f"ประวัติการศึกษา:\n{st.session_state.edu_history}")
    if st.session_state.scholarship.strip(): bio_parts.append(f"ทุนการศึกษา:\n{st.session_state.scholarship}")
    if st.session_state.publications.strip(): bio_parts.append(f"ผลงานทางวิชาการ:\n{st.session_state.publications}")
    thesis.author_biography = "\n\n".join(bio_parts)

st.divider()
st.subheader("🛠️ ขั้นตอนสุดท้าย: ออกเล่มวิทยานิพนธ์")

try:
    docx_buffer = build_docx_file(thesis, citation_style)
    st.download_button(
        label="📥 ดาวน์โหลดไฟล์วิทยานิพนธ์ฉบับจัดหน้าสำเร็จ (.docx)",
        data=docx_buffer,
        file_name=f"Thesis_{thesis.author_en.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary",
        use_container_width=True
    )
except Exception as e:
    st.error(f"⚠️ เกิดข้อผิดพลาดในการสร้างไฟล์: {e}")